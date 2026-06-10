from pathlib import Path
import json
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


def zh(text):
    return text.encode("utf-8").decode("utf-8")


ROOT = Path(r"D:\TestReport")
STATE_PATH = ROOT / "app-state.json"
OUT_DIR = ROOT / "output" / "doc"
OUT_DIR.mkdir(parents=True, exist_ok=True)
OUT_PATH = OUT_DIR / "test-report-template-current.docx"


with STATE_PATH.open("r", encoding="utf-8") as f:
    payload = json.load(f)
state = payload.get("state", {})

batches = state.get("batches", [])
tasks = state.get("tasks", [])
cases = state.get("cases", [])
bugs = state.get("bugs", [])
report_conclusion = state.get("reportConclusion") or zh("暂无补充结论。")
last_generation = state.get("lastGeneration") or {}

batch = batches[0] if batches else {}
task = tasks[0] if tasks else {}
module_name = task.get("moduleName") or batch.get("moduleName") or (cases[0].get("module") if cases else zh("未选择"))
batch_version = task.get("batchVersion") or batch.get("version") or (cases[0].get("batchVersion") if cases else zh("未选择"))
task_name = task.get("name") or zh("未选择")
owners = task.get("owners") or batch.get("owners") or []
if isinstance(owners, str):
    owners = [x.strip() for x in owners.replace("，", "、").split("、") if x.strip()]
owner_text = "、".join(owners) if owners else zh("未分配")

status_order = [zh("未执行"), zh("通过"), zh("失败"), zh("阻塞")]
status_counts = {k: 0 for k in status_order}
for item in cases:
    key = item.get("executionStatus") or zh("未执行")
    status_counts[key] = status_counts.get(key, 0) + 1


def count_bugs(items, key, order):
    result = {k: 0 for k in order}
    for item in items:
        current = item.get(key) or order[0]
        result[current] = result.get(current, 0) + 1
    return result


bug_status_order = [zh("新建"), zh("已提交"), zh("已修复"), zh("已验证"), zh("已关闭")]
bug_severity_order = [zh("严重"), zh("中"), zh("低")]
bug_status_counts = count_bugs(bugs, "status", bug_status_order)
bug_severity_counts = count_bugs(bugs, "severity", bug_severity_order)
failed_cases = [item for item in cases if item.get("executionStatus") == zh("失败")]
unresolved_bugs = [item for item in bugs if item.get("status") not in [zh("已修复"), zh("已验证"), zh("已关闭")]]
open_bugs = len([item for item in bugs if item.get("status") not in [zh("已验证"), zh("已关闭")]])
passed = status_counts.get(zh("通过"), 0)
executed = len(cases) - status_counts.get(zh("未执行"), 0)
pass_rate = f"{round((passed / executed) * 100)}%" if executed else "0%"
source_type = zh("API内容") if last_generation.get("type") == "api" else zh("需求内容")

if bug_severity_counts.get(zh("严重"), 0) > 0 or status_counts.get(zh("失败"), 0) > 0:
    decision = (zh("有风险"), zh("存在失败用例或严重 BUG，建议修复后再回归。"))
elif status_counts.get(zh("阻塞"), 0) > 0 or open_bugs > 0:
    decision = (zh("需关注"), zh("当前还有阻塞项或待跟进 BUG，上线前建议继续确认。"))
else:
    decision = (zh("可发布"), zh("当前执行结果稳定，未发现明显发布阻塞。"))


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_run_font(run, size, bold=False, color=None):
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    if color:
        run.font.color.rgb = RGBColor(*color)


def set_cell_text(cell, text, bold=False, size=10.5):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(str(text))
    set_run_font(run, size, bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_heading(doc, text, size=14):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    set_run_font(run, size, bold=True)


def add_kv_table(doc, rows):
    table = doc.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for key, value in rows:
        cells = table.add_row().cells
        set_cell_text(cells[0], key, bold=True)
        set_cell_text(cells[1], value)
        set_cell_shading(cells[0], "F3F6FB")
    return table


def add_matrix_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for idx, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[idx], header, bold=True, size=10)
        set_cell_shading(table.rows[0].cells[idx], "EAF0F8")
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_text(cells[idx], value, size=10)
    return table


doc = Document()
section = doc.sections[0]
section.top_margin = Cm(2)
section.bottom_margin = Cm(2)
section.left_margin = Cm(2)
section.right_margin = Cm(2)

style = doc.styles["Normal"]
style.font.name = "Microsoft YaHei"
style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
style.font.size = Pt(10.5)

p = doc.add_paragraph()
p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
run = p.add_run(zh("测试报告"))
set_run_font(run, 20, bold=True)

sub = doc.add_paragraph()
sub.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
sub_run = sub.add_run(task_name if task_name != zh("未选择") else zh("当前测试范围"))
set_run_font(sub_run, 12, color=(95, 107, 122))

doc.add_paragraph("")

sections = [
    (
        zh("文档信息"),
        [
            (zh("报告名称"), zh("测试报告")),
            (zh("版本号"), batch_version),
            (zh("测试任务"), task_name),
            (zh("测试负责人"), owner_text),
            (zh("生成方式"), zh("系统当前报告导出为 DOCX 模板")),
            (zh("报告范围"), f"{zh('业务')}：{module_name} / {zh('版本')}：{batch_version} / {zh('任务')}：{task_name}"),
        ],
    ),
    (
        zh("测试范围摘要"),
        [
            (zh("测试业务"), module_name),
            (zh("来源类型"), source_type),
            (zh("测试内容"), task.get("scope") or batch.get("scope") or zh("未填写")),
            (zh("测试对象"), task_name if task_name != zh("未选择") else batch_version),
            (zh("当前结论"), decision[0]),
        ],
    ),
    (
        zh("用例执行统计"),
        [
            (zh("测试用例总数"), len(cases)),
            (zh("执行用例数"), executed),
            (zh("成功用例数"), passed),
            (zh("失败用例数"), status_counts.get(zh("失败"), 0)),
            (zh("阻塞用例数"), status_counts.get(zh("阻塞"), 0)),
            (zh("未执行用例数"), status_counts.get(zh("未执行"), 0)),
            (zh("通过率"), pass_rate),
        ],
    ),
    (
        zh("缺陷统计"),
        [
            (zh("BUG总数"), len(bugs)),
            (zh("待跟进BUG"), open_bugs),
            (zh("新建"), bug_status_counts.get(zh("新建"), 0)),
            (zh("已提交"), bug_status_counts.get(zh("已提交"), 0)),
            (zh("已修复"), bug_status_counts.get(zh("已修复"), 0)),
            (zh("已验证"), bug_status_counts.get(zh("已验证"), 0)),
            (zh("已关闭"), bug_status_counts.get(zh("已关闭"), 0)),
            (zh("严重"), bug_severity_counts.get(zh("严重"), 0)),
            (zh("中"), bug_severity_counts.get(zh("中"), 0)),
            (zh("低"), bug_severity_counts.get(zh("低"), 0)),
        ],
    ),
    (
        zh("风险与结论"),
        [
            (zh("发布建议"), decision[0]),
            (zh("结论说明"), decision[1]),
            (zh("补充结论"), report_conclusion),
        ],
    ),
    (
        zh("测试结论与建议"),
        [
            (zh("当前判断"), decision[0]),
            (
                zh("建议动作"),
                zh("可以进入发布确认，保留最终抽查记录。")
                if decision[0] == zh("可发布")
                else zh("建议修复问题后补充回归，再更新本报告。"),
            ),
            (zh("失败用例"), f"{status_counts.get(zh('失败'), 0)} {zh('条')}"),
            (zh("阻塞用例"), f"{status_counts.get(zh('阻塞'), 0)} {zh('条')}"),
            (zh("待跟进BUG"), f"{open_bugs} {zh('个')}"),
            (zh("未执行用例"), f"{status_counts.get(zh('未执行'), 0)} {zh('条')}"),
        ],
    ),
]

for title, rows in sections:
    add_heading(doc, title, size=14)
    add_kv_table(doc, rows)
    doc.add_paragraph("")

add_heading(doc, zh("重点关注"), size=14)
add_heading(doc, zh("失败的用例"), size=12)
failed_rows = [
    [
        idx,
        item.get("title") or "",
        item.get("taskName") or zh("未分任务"),
        item.get("batchVersion") or batch_version,
        item.get("module") or module_name,
        zh("失败"),
    ]
    for idx, item in enumerate(failed_cases, 1)
]
if not failed_rows:
    failed_rows = [["-", zh("当前没有失败用例"), "", "", "", ""]]
add_matrix_table(doc, [zh("序号"), zh("用例标题"), zh("任务"), zh("版本"), zh("业务"), zh("状态")], failed_rows)

doc.add_paragraph("")
add_heading(doc, zh("未修复的BUG"), size=12)
bug_rows = [
    [
        idx,
        item.get("title") or "",
        item.get("status") or "",
        item.get("severity") or "",
        item.get("taskName") or zh("未分任务"),
        item.get("owner") or zh("未填写"),
        item.get("link") or zh("未填写"),
        item.get("note") or zh("暂无补充说明"),
    ]
    for idx, item in enumerate(unresolved_bugs, 1)
]
if not bug_rows:
    bug_rows = [["-", zh("当前没有未修复BUG"), "", "", "", "", "", ""]]
add_matrix_table(
    doc,
    [zh("序号"), zh("BUG标题"), zh("状态"), zh("严重级别"), zh("任务"), zh("负责人"), zh("Lark链接"), zh("详情")],
    bug_rows,
)

doc.save(OUT_PATH)
print(OUT_PATH)
