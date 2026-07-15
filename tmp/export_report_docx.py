from pathlib import Path
import json
import sys

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


COLOR_RED = RGBColor(190, 55, 65)
COLOR_ORANGE = RGBColor(174, 112, 18)
COLOR_GREEN = RGBColor(28, 126, 79)
COLOR_GRAY = RGBColor(92, 108, 116)

FILL_RED = "FFF0F1"
FILL_ORANGE = "FFF5E3"
FILL_GREEN = "EAF8EF"
FILL_GRAY = "F1F5F6"
FILL_HEADER = "E7F3F4"
FILL_LABEL = "F4F7F8"
FILL_ROW_RISK = "FFF7F7"
FILL_ROW_WARN = "FFF9F0"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    tc_pr.append(shading)


def set_run_font(run, size, bold=False, color=None):
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    if color:
        run.font.color.rgb = color


def get_status_style(value):
    text = str(value or "").strip()
    if text in ["有风险", "失败", "新建", "严重", "P0"]:
        return COLOR_RED, FILL_RED
    if text in ["需关注", "阻塞", "已提交", "中", "P1", "待回归"]:
        return COLOR_ORANGE, FILL_ORANGE
    if text in ["可发布", "通过", "已修复", "已验证", "已关闭", "低", "P2"]:
        return COLOR_GREEN, FILL_GREEN
    return COLOR_GRAY, FILL_GRAY


def set_cell_text(cell, text, bold=False, size=10.5, color=None, fill=None):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(str(text if text is not None else ""))
    set_run_font(run, size, bold=bold, color=color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    if fill:
        set_cell_shading(cell, fill)


def add_heading(doc, text, size=14):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(8)
    paragraph.paragraph_format.space_after = Pt(6)
    run = paragraph.add_run(text)
    set_run_font(run, size, bold=True)


def add_kv_table(doc, rows):
    table = doc.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for key, value, *options in rows:
        highlight = bool(options[0]) if options else False
        cells = table.add_row().cells
        set_cell_text(cells[0], key, bold=True, fill=FILL_LABEL)
        if highlight:
            color, fill = get_status_style(value)
            set_cell_text(cells[1], value, color=color, fill=fill)
        else:
            set_cell_text(cells[1], value)
    return table


def add_matrix_table(doc, headers, rows, status_columns=None, row_fill_fn=None):
    status_columns = status_columns or set()
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for index, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[index], header, bold=True, size=9.5, fill=FILL_HEADER)
    for row in rows:
        cells = table.add_row().cells
        row_fill = row_fill_fn(row) if row_fill_fn else None
        for index, value in enumerate(row):
            if index in status_columns:
                color, fill = get_status_style(value)
                set_cell_text(cells[index], value, size=9.5, color=color, fill=fill)
            else:
                set_cell_text(cells[index], value, size=9.5, fill=row_fill)
    return table


def bug_row_fill(row):
    status = str(row[2] or "").strip()
    severity = str(row[3] or "").strip()
    if severity == "严重" or status == "新建":
        return FILL_ROW_RISK
    if status in ["已提交", "待回归"] or severity == "中":
        return FILL_ROW_WARN
    return None


def add_document_header(doc, report):
    title = doc.add_paragraph()
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    set_run_font(title.add_run("测试报告"), 20, bold=True)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    set_run_font(subtitle.add_run(report.get("heroTitle") or "测试结果汇总"), 12, color=COLOR_GRAY)
    doc.add_paragraph("")


def build_conclusion_rows(items, fallback):
    source = items if isinstance(items, list) else []
    if source and all(isinstance(item, str) for item in source):
        source = [source[index:index + 2] for index in range(0, len(source), 2)]

    rows = []
    for item in source:
        if isinstance(item, (list, tuple)) and len(item) >= 2:
            label, detail = item[0], item[1]
        elif isinstance(item, dict):
            label = item.get("label") or item.get("title") or "建议"
            detail = item.get("detail") or item.get("value") or item.get("description") or ""
        else:
            continue
        rows.append([label, detail, label == "当前判断"])

    return rows or [["当前判断", fallback or "暂无补充建议。", True]]


def build_sections(report, conclusion):
    release = report.get("releaseDecision", {})
    status_counts = report.get("statusCounts", {})
    bug_status_counts = report.get("bugStatusCounts", {})
    bug_severity_counts = report.get("bugSeverityCounts", {})
    scoped_bugs = report.get("scope", {}).get("bugs", [])

    return [
        (
            "文档信息",
            [
                ["报告名称", "测试报告"],
                ["版本号", report.get("batchVersion", "")],
                ["测试任务", report.get("taskName", "")],
                ["生成时间", report.get("generatedAt", "")],
                ["当前结论", release.get("label", ""), True],
            ],
        ),
        (
            "用例执行统计",
            [
                ["测试用例总数", report.get("total", 0)],
                ["执行用例数", report.get("executed", 0)],
                ["成功用例数", report.get("passed", 0), True],
                ["失败用例数", status_counts.get("失败", 0), True],
                ["阻塞用例数", status_counts.get("阻塞", 0), True],
                ["未执行用例数", status_counts.get("未执行", 0)],
                ["通过率", report.get("passRate", "0%"), True],
            ],
        ),
        (
            "缺陷统计",
            [
                ["BUG总数", len(scoped_bugs)],
                ["待跟进BUG", report.get("openBugs", 0), True],
                ["失败用例对应BUG数", report.get("failedCaseBugCount", 0), True],
                ["新建", bug_status_counts.get("新建", 0), True],
                ["已提交", bug_status_counts.get("已提交", 0), True],
                ["已修复", bug_status_counts.get("已修复", 0), True],
                ["待回归", bug_status_counts.get("待回归", 0), True],
                ["已验证", bug_status_counts.get("已验证", 0), True],
                ["已关闭", bug_status_counts.get("已关闭", 0), True],
                ["严重", bug_severity_counts.get("严重", 0), True],
                ["中", bug_severity_counts.get("中", 0), True],
                ["低", bug_severity_counts.get("低", 0), True],
            ],
        ),
        (
            "风险与结论",
            [
                ["发布建议", release.get("label", ""), True],
                ["结论说明", release.get("desc", "")],
                ["补充结论", conclusion],
            ],
        ),
        (
            "阻塞原因汇总",
            [["阻塞说明", report.get("blockedSummary", "当前没有阻塞用例。")]],
        ),
        (
            "测试结论与建议",
            build_conclusion_rows(
                report.get("conclusionAdviceItems", []),
                release.get("desc", "暂无补充建议。"),
            ),
        ),
    ]


def add_focus_section(doc, report):
    add_heading(doc, "重点关注", size=14)
    add_heading(doc, "失败的用例", size=12)
    failed_cases = report.get("failedCases", [])
    failed_rows = [
        [
            index,
            item.get("title") or "",
            item.get("taskName") or "未分任务",
            item.get("priority") or "",
            "失败",
            len([bug for bug in report.get("unresolvedBugs", []) if bug.get("caseId") == item.get("id")]),
        ]
        for index, item in enumerate(failed_cases, 1)
    ] or [["-", "当前没有失败用例", "", "", "", ""]]
    add_matrix_table(
        doc,
        ["序号", "用例标题", "任务", "优先级", "状态", "关联BUG数"],
        failed_rows,
        status_columns={3, 4},
        row_fill_fn=(lambda _row: FILL_ROW_RISK) if failed_cases else None,
    )

    doc.add_paragraph("")
    add_heading(doc, "未关闭的BUG", size=12)
    unresolved_bugs = report.get("unresolvedBugs", [])
    bug_rows = [
        [
            index,
            item.get("title") or "",
            item.get("status") or "",
            item.get("severity") or "",
            item.get("taskName") or "未分任务",
            item.get("note") or "暂无补充说明",
        ]
        for index, item in enumerate(unresolved_bugs, 1)
    ] or [["-", "当前没有未关闭BUG", "", "", "", ""]]
    add_matrix_table(
        doc,
        ["序号", "BUG标题", "状态", "严重级别", "任务", "详情"],
        bug_rows,
        status_columns={2, 3},
        row_fill_fn=bug_row_fill if unresolved_bugs else None,
    )


def main():
    payload_path = Path(sys.argv[1])
    payload = json.loads(payload_path.read_text(encoding="utf-8"))
    report = payload["report"]
    conclusion = payload.get("reportConclusion") or "暂无补充结论。"
    output_path = Path(payload["outputPath"])
    output_path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)

    normal = doc.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)

    add_document_header(doc, report)
    for heading, rows in build_sections(report, conclusion):
        add_heading(doc, heading, size=14)
        add_kv_table(doc, rows)
        doc.add_paragraph("")
    add_focus_section(doc, report)

    doc.save(output_path)
    print(str(output_path))


if __name__ == "__main__":
    main()
